import tkinter as tk
from tkinter.scrolledtext import ScrolledText
from os.path import join
import docx

class Application():

    TOP_SPACE, LEFT_SPACE = 15, 30
    TEXT_SIZE = 9
    UPPERCASE, LOWERCASE = None, None

    def __init__(self, master=None):
        self.master = master
        self.build_gui()

    def build_gui(self, width=20, height=30, cols=4):
        # Headers
        self.headers = []
        for h in range(4):
            e = tk.Entry(self.master, width=width+10)
            e.grid(row=0, column=h)
            self.headers.append(e)
            e = None
        
        # Rows
        self.rows = []
        for t in range(4):
            r = ScrolledText(self.master, width=width, height=height)
            r.grid(row=1, column=t)
            self.rows.append(r)
            r = None
        
        # Checkboxes
        self.HEADER_UPPERCASE = tk.IntVar()
        self.HEADER_LOWERCASE = tk.IntVar()
        tk.Checkbutton(
            self.master, 
            text="en MAJUSCULES", 
            variable=self.HEADER_UPPERCASE
        ).grid(row=0, column=cols, sticky='N')
        tk.Checkbutton(
            self.master, 
            text="en minuscules", 
            variable=self.HEADER_LOWERCASE
        ).grid(row=0, column=cols+1, sticky='N')
        
        self.ROWS_UPPERCASE = tk.IntVar()
        self.ROWS_LOWERCASE = tk.IntVar()
        tk.Checkbutton(
            self.master, 
            text="en MAJUSCULES", 
            variable=self.ROWS_UPPERCASE
        ).grid(row=1, column=cols, sticky='N')
        tk.Checkbutton(
            self.master, 
            text="en minuscules", 
            variable=self.ROWS_LOWERCASE
        ).grid(row=1, column=cols+1, sticky='N')
        
        # Button
        self.do = tk.Button(text='Generar etiquetes', command=self.do_action)
        self.do.grid(row=2, column=cols+1, sticky='EW')
        
    
    def do_action(self):
        self.generate_excel(
            data = self.prepare_data(), 
            entry_name = join(
                join('.','input'),
                'Etiquetes.docx'
            ),
            output_name = join(
                join('.', 'output'),
                'Result.docx'
            )
        )
        self.master.destroy()
    
    
    def format_cell(self, header, col):
        replacements = [
            ('\t', ' ')
            #, ('  ',' ')
        ]
    
        header = header.strip()
        if self.HEADER_UPPERCASE.get(): header = header.upper()
        if self.HEADER_LOWERCASE.get(): header = header.lower()
        
        col = col.strip()
        for search, replace in replacements:
            col = col.replace(search, replace)
        if self.ROWS_UPPERCASE.get(): col = col.upper()
        if self.ROWS_LOWERCASE.get(): col = col.lower()
        
        return "{header}{sep}{text}{newline}" \
                        .format(
                            header=header,
                            text=col,
                            sep=": " if header and col else "",
                            newline="\n"
                        )
    
    
    def prepare_data(self):
        """
        Prepares cell values as shown below:
        
        header1   header2
        -------   -------
        row1_1    row_1_2      
        row2_1    row_2_2
        
        cell[0] = 'header1: row1_1\nheader2: row_1_2'
        cell[1] = 'header2: row2_1\nheader2: row_2_2'
        """
        headers_text = [ h.get() for h in self.headers ]
        rows_text = [ r.get(1.0, tk.END) for r in self.rows ]
        
        cells = [""]*(len(headers_text) * len(rows_text[0].splitlines()))
        for header, col in zip(headers_text, rows_text):
            for i in range(len(col.splitlines())):
                if not col=="":
                    cells[i] += self.format_cell(header, col.splitlines()[i])
        return [cell for cell in cells if cell]

        
    def generate_excel(self, data, entry_name, output_name):
        document = docx.Document(entry_name)
        
        i=0
        for row_i, row in enumerate(document.tables[0].rows):
            for col_i, cell in enumerate(row.cells):
                # add text
                if i<len(data):
                    cell.text = data[i]
                    i+=1
                elif col_i==0:
                    # TODO: remove row or cells
                    pass
                
                # format cell
                for paragraph in cell.paragraphs:
                    p_format = paragraph.paragraph_format
                    p_format.left_indent = docx.shared.Pt(self.LEFT_SPACE)
                    p_format.space_before = docx.shared.Pt(self.TOP_SPACE)
                    
                    for run in paragraph.runs:
                        font = run.font
                        font.size = docx.shared.Pt(self.TEXT_SIZE)

        document.save(output_name)

if __name__ == '__main__':
    root = tk.Tk()
    app = Application(root)
    root.mainloop()


